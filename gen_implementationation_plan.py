"""
生成《项目总体实施方案》文档
基于模板 5 CWRF区域气候模式重点区域和流域尺度实施方案.docx 的结构
内容映射为：基于深度学习的对流云识别与外推算项目
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATE = r"files\样例\5 CWRF区域气候模式重点区域和流域尺度实施方案.docx"
OUTPUT = r"files\样例\5_实施方案_填充版.docx"


def set_paragraph_font(paragraph, font_name='宋体', font_size=12, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.font.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def main():
    doc = Document(TEMPLATE)
    # 清除模板中已有的内容段落（保留文档结构）
    # 策略：遍历所有段落，在对应标题后插入内容

    # 由于模板已有标题结构，我们直接在模板基础上填充内容
    # 找到每个标题段落并在其后添加内容

    # 获取所有段落的索引和样式信息
    paragraphs = doc.paragraphs
    heading_indices = {}
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith('Heading'):
            heading_indices[p.text.strip()] = i

    # 由于python-docx不支持在指定位置插入，我们采用新策略：
    # 创建一个新文档，按模板结构重新构建

    new_doc = Document()

    # 设置默认字体
    style = new_doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ========== 第一部分：项目总体实施方案 ==========
    new_doc.add_heading('项目总体实施方案', level=1)

    new_doc.add_heading('项目背景', level=2)
    add_para(new_doc,
        '强对流天气是气象防灾减灾的重点关注对象，其对流云团的快速生消演变对人民生命财产安全构成严重威胁。'
        '风云四号（FY-4）气象卫星搭载的先进静止轨道辐射成像仪（AGRI）具有高时间分辨率、多光谱通道观测优势，'
        '为对流云的精细化监测和临近预警提供了重要的数据支撑。', indent=True)
    add_para(new_doc,
        '本项目旨在基于深度学习方法，利用FY-4 AGRI多通道亮温数据，实现覆盖范围10°-40°N、110°-145°E区域内'
        '对流云团的智能识别与1小时外推算（nowcasting）。项目研发两套核心深度学习模型：'
        '（1）DIF-UNet双流识别融合网络，用于对流云高精度识别；'
        '（2）SatExNet卫星回波外推网络，用于对流云未来1小时演变预测。'
        '项目成果将直接服务于气象业务部门对流天气的实时监测与短临预警。', indent=True)

    new_doc.add_heading('项目实施原则', level=2)
    add_para(new_doc, '（1）科学性与先进性原则：采用国际前沿的深度学习架构（Swin Transformer、U-Net等），确保技术方案处于领域领先水平。', indent=True)
    add_para(new_doc, '（2）实用性原则：系统设计充分考虑气象业务实际需求，模型推理速度满足实时处理要求，输出产品可直接对接业务平台。', indent=True)
    add_para(new_doc, '（3）可扩展性原则：代码架构采用模块化设计，支持后续增加新的卫星通道、扩展预报时效、适配其他区域。', indent=True)
    add_para(new_doc, '（4）可靠性原则：建立完善的训练、验证、测试流程，确保模型在不同季节、不同天气条件下均具有稳定的识别与预报能力。', indent=True)
    add_para(new_doc, '（5）规范性原则：严格遵循软件工程管理规范，保证代码质量、文档完整性和过程可追溯。', indent=True)

    new_doc.add_heading('项目总体推进计划', level=2)
    new_doc.add_heading('项目总体目标', level=3)
    add_para(new_doc,
        '研发基于FY-4 AGRI卫星数据的对流云智能识别与外推算系统，实现对流云识别准确率≥85%、'
        '1小时外推算相关系数≥0.7的核心指标，交付完整的源代码、技术文档和模型文件。', indent=True)

    new_doc.add_heading('业务目标', level=3)
    add_para(new_doc, '（1）构建覆盖10°-40°N、110°-145°E区域、空间分辨率4km的对流云识别数据集，标注样本不少于2920个时次。', indent=True)
    add_para(new_doc, '（2）研发DIF-UNet双流识别融合网络，融合Swin Transformer全局特征与ResNet局部特征，实现多光谱通道对流云精准识别。', indent=True)
    add_para(new_doc, '（3）研发SatExNet卫星回波外推网络，基于自回归架构实现n-1到n+1小时的逐小时对流云演变预测。', indent=True)
    add_para(new_doc, '（4）建立完整的模型训练、评估、推理流水线，支持业务化运行。', indent=True)

    new_doc.add_heading('项目建设周期与时间进度要求', level=3)
    add_para(new_doc, '项目总建设周期为12个月，分为以下阶段：', indent=True)
    add_para(new_doc, '第一阶段（第1-2月）：需求分析与数据准备。完成FY-4 AGRI数据收集、预处理、质控和标注，构建训练数据集。', indent=True)
    add_para(new_doc, '第二阶段（第3-5月）：算法设计与模型开发。完成DIF-UNet和SatExNet两大核心模型的架构设计、编码实现和初步训练。', indent=True)
    add_para(new_doc, '第三阶段（第6-8月）：模型优化与实验验证。开展超参数调优、消融实验、对比实验，完成模型性能评估。', indent=True)
    add_para(new_doc, '第四阶段（第9-10月）：系统集成与测试。完成推理流水线开发、系统联调测试和性能优化。', indent=True)
    add_para(new_doc, '第五阶段（第11-12月）：文档编写与项目验收。完成技术报告、实施方案、用户手册等文档编写，组织项目验收。', indent=True)

    new_doc.add_heading('项目范围', level=2)
    add_para(new_doc,
        '本项目范围包括：（1）FY-4 AGRI多通道亮温数据的预处理与数据集构建；'
        '（2）DIF-UNet双流识别融合网络的设计、训练与评估；'
        '（3）SatExNet卫星回波外推网络的设计、训练与评估；'
        '（4）模型推理服务与业务化接口开发；'
        '（5）技术文档、测试报告、用户手册编写；'
        '（6）项目培训与技术支持。', indent=True)

    new_doc.add_heading('系统实施过程的质量保证活动说明', level=2)

    new_doc.add_heading('项目启动', level=3)
    add_para(new_doc, '项目启动阶段主要完成以下工作：成立项目组，明确各成员职责分工；召开项目启动会，确认项目目标、范围、进度计划；建立项目沟通机制和配置管理环境。', indent=True)

    new_doc.add_heading('需求分析阶段', level=3)
    add_para(new_doc, '需求分析阶段主要完成：与气象业务部门对接，明确对流云识别与外推算的具体业务需求；分析FY-4 AGRI数据格式（HDF5/NetCDF）、通道特性、时空分辨率；确定模型输入输出规格和性能指标要求；编写需求规格说明书。', indent=True)

    new_doc.add_heading('设计开发阶段', level=3)
    new_doc.add_heading('概要设计', level=4)
    add_para(new_doc, '概要设计包括：系统总体架构设计（数据层、算法层、服务层）；DIF-UNet网络架构设计（双流编码器、SIFM融合模块、CLFIM多尺度模块）；SatExNet网络架构设计（自回归骨干网络、时序编码模块）；训练策略设计（损失函数、优化器、学习率调度）。', indent=True)

    new_doc.add_heading('实施方案', level=4)
    add_para(new_doc, '实施方案详细规定各模块的开发计划、接口定义、数据流转方式。DIF-UNet模块采用PyTorch框架实现，使用Swin Transformer（预训练于ImageNet）和ResNet-50作为双流编码器，通过SIFM模块进行特征融合，CLFIM模块实现多尺度上下文聚合。SatExNet模块采用自回归架构，利用历史多时次输入预测未来对流云演变。', indent=True)

    new_doc.add_heading('程序编码', level=4)
    add_para(new_doc, '程序编码遵循PEP 8规范，采用模块化设计。主要代码模块包括：数据加载模块（data_loader.py）、模型定义模块（models/dif_unet.py、models/satexnet.py）、损失函数模块（losses.py）、训练脚本（train_dif_unet.py、train_satexnet.py）、推理模块（inference.py）。所有代码使用Git进行版本管理。', indent=True)

    new_doc.add_heading('系统测试阶段', level=3)
    new_doc.add_heading('分模块测试', level=4)
    add_para(new_doc, '对各独立模块进行单元测试：数据加载模块测试数据格式解析、归一化、增强的正确性；模型模块测试前向传播的维度匹配、梯度流通性；损失函数模块测试数值稳定性和梯度计算。', indent=True)

    new_doc.add_heading('结合测试', level=4)
    add_para(new_doc, '进行端到端集成测试：验证从数据输入到模型输出的完整流水线；测试训练过程的收敛性和稳定性；验证推理模块的输出格式和性能指标计算的正确性。', indent=True)

    new_doc.add_heading('质量控制体系', level=3)
    new_doc.add_heading('开发品质标准', level=4)
    add_para(new_doc, '代码覆盖率不低于80%；所有公共函数必须有文档字符串；关键算法必须有单元测试；代码审查通过后方可合并至主分支。', indent=True)

    new_doc.add_heading('测试密度与bug检出率', level=4)
    add_para(new_doc, '每千行代码缺陷数不超过5个；严重缺陷修复率100%；一般缺陷修复率不低于95%。', indent=True)

    new_doc.add_heading('问题处置能力标准', level=4)
    add_para(new_doc, '关键问题24小时内响应，48小时内提出解决方案；一般问题3个工作日内解决。', indent=True)

    new_doc.add_heading('任务完成度的进度率标准', level=4)
    add_para(new_doc, '各里程碑任务完成率不低于90%；关键路径任务完成率100%。', indent=True)

    new_doc.add_heading('正规化编程生产性标准', level=4)
    add_para(new_doc, '遵循Python PEP 8编码规范；使用类型注解提高代码可读性；采用统一的日志记录框架；配置文件与代码分离。', indent=True)

    new_doc.add_heading('系统交付与用户测试', level=3)
    add_para(new_doc, '系统交付前进行用户验收测试（UAT），由气象业务人员使用实际业务数据进行测试，验证模型在真实业务场景下的表现。用户测试重点包括：识别准确率、外推算精度、推理速度、系统稳定性。', indent=True)

    new_doc.add_heading('实施过程提交文件汇总', level=3)
    add_para(new_doc, '项目实施过程中需提交的文件包括：需求规格说明书、概要设计说明书、详细设计说明书、测试计划与测试报告、用户操作手册、技术报告、源代码及说明文档、模型权重文件。', indent=True)

    # ========== 第二部分：项目管理方案 ==========
    new_doc.add_heading('项目管理方案', level=1)

    new_doc.add_heading('项目管理组织结构', level=2)
    new_doc.add_heading('项目各方角色与责任', level=3)
    add_para(new_doc, '项目设立项目管理委员会、技术专家组和项目执行团队三级组织架构。项目管理委员会负责项目整体决策和资源协调；技术专家组负责技术方案评审和质量把关；项目执行团队负责具体研发工作实施。', indent=True)
    add_para(new_doc, '项目执行团队设置以下岗位：项目经理1名，负责总体协调和进度管控；算法工程师2-3名，负责DIF-UNet和SatExNet模型的设计与实现；数据工程师1名，负责FY-4数据处理和数据集构建；测试工程师1名，负责系统测试和质量保障。', indent=True)

    new_doc.add_heading('用户角色与责任', level=3)
    add_para(new_doc, '用户方（气象业务部门）负责：提供业务需求说明和验收标准；提供历史观测数据用于模型验证；参与阶段性评审和验收测试；提供部署环境和业务对接支持。', indent=True)

    new_doc.add_heading('项目范围管理', level=2)
    add_para(new_doc, '采用需求变更控制流程管理项目范围。任何范围变更需提交变更申请，经项目管理委员会评估影响后审批。范围基线在需求分析阶段结束后建立。', indent=True)

    new_doc.add_heading('项目进度管理', level=2)
    add_para(new_doc, '使用甘特图进行进度管理，每两周召开项目进度会议，跟踪各任务完成情况。关键里程碑设置检查点，对偏差超过10%的任务启动纠偏措施。', indent=True)

    new_doc.add_heading('项目风险管理', level=2)
    new_doc.add_heading('技术风险', level=3)
    add_para(new_doc, '主要技术风险包括：（1）模型训练不收敛或性能不达标——应对措施：预留充分的调参时间，准备多种备选架构方案；（2）GPU计算资源不足——应对措施：采用混合精度训练、梯度累积等技术提高计算效率；（3）数据质量不满足要求——应对措施：建立严格的数据质控流程，准备数据增强策略。', indent=True)

    new_doc.add_heading('需求风险', level=3)
    add_para(new_doc, '需求变更可能导致开发返工。应对措施：在需求分析阶段充分沟通确认，建立变更控制流程，预留10%的进度缓冲。', indent=True)

    new_doc.add_heading('协调与沟通风险', level=3)
    add_para(new_doc, '项目组与用户方沟通不畅可能导致需求理解偏差。应对措施：建立定期沟通机制，每两周向用户方汇报进展，关键决策点邀请用户方参与评审。', indent=True)

    new_doc.add_heading('项目人员风险', level=3)
    add_para(new_doc, '核心技术人员变动可能影响项目进度。应对措施：关键岗位设置AB角，做好技术文档和知识传承。', indent=True)

    new_doc.add_heading('质量管理计划', level=2)
    new_doc.add_heading('质量管理体系标准', level=3)
    add_para(new_doc, '项目遵循ISO 9001质量管理体系标准，结合CMMI 3级过程能力要求，建立项目质量管理制度。', indent=True)

    new_doc.add_heading('质量控制过程', level=3)
    add_para(new_doc, '质量控制贯穿项目全生命周期，包括：需求评审、设计评审、代码审查、测试验证、用户验收五个质量控制环节。', indent=True)

    new_doc.add_heading('质量评定计划', level=3)
    add_para(new_doc, '每月进行一次质量评定，评估代码质量、测试覆盖率、缺陷密度等指标，形成质量报告。', indent=True)

    new_doc.add_heading('质量管理措施', level=3)
    add_para(new_doc, '建立代码审查制度，所有代码提交必须经过至少一名其他工程师审查通过；使用自动化测试工具持续验证代码正确性；定期进行代码质量扫描（使用pylint、mypy等工具）。', indent=True)

    new_doc.add_heading('软件质量控制', level=3)
    new_doc.add_heading('阶段性评审', level=4)
    add_para(new_doc, '在需求分析、概要设计、详细设计、编码完成、测试完成等关键阶段设置评审点，由技术专家组进行评审，评审通过后方可进入下一阶段。', indent=True)

    new_doc.add_heading('测试', level=4)
    add_para(new_doc, '建立三级测试体系：单元测试（开发人员执行，覆盖率≥80%）、集成测试（测试团队执行）、系统测试（独立测试团队执行）。', indent=True)

    new_doc.add_heading('项目协调与合作计划', level=2)
    new_doc.add_heading('协调与合作管理方案', level=3)
    add_para(new_doc, '建立项目组内部协作机制和用户方合作机制。内部采用每日站会、每周周报制度；外部采用双周汇报、里程碑评审制度。', indent=True)

    new_doc.add_heading('协调手段', level=3)
    add_para(new_doc, '使用项目管理工具（如GitLab、Jira）进行任务分配和进度跟踪；使用即时通讯工具进行日常沟通；使用文档协作平台进行技术文档共享。', indent=True)

    new_doc.add_heading('配置管理', level=2)
    new_doc.add_heading('配置管理和版本控制', level=3)
    add_para(new_doc, '使用Git进行代码版本管理，采用Git Flow分支策略。主分支（main）保持稳定，开发分支（develop）用于日常开发，功能分支（feature）用于新功能开发。模型权重文件使用DVC（Data Version Control）进行管理。', indent=True)

    new_doc.add_heading('变更管理的方法', level=3)
    add_para(new_doc, '所有变更通过变更请求流程管理：提交变更申请→影响评估→审批→实施→验证→关闭。紧急变更可先行实施，但需在24小时内补充变更申请。', indent=True)

    new_doc.add_heading('文档管理', level=2)
    add_para(new_doc, '项目文档采用统一编号和版本管理。文档类型包括：管理类文档（项目计划、进度报告）、技术类文档（设计文档、测试报告）、交付类文档（用户手册、培训材料）。所有文档存储于项目文档库，定期备份。', indent=True)

    new_doc.add_heading('人员管理', level=2)
    add_para(new_doc, '建立人员管理制度，明确岗位职责、绩效考核标准和培训发展计划。定期组织技术分享和培训活动，提升团队整体技术水平。', indent=True)

    new_doc.add_heading('保密管理', level=2)
    add_para(new_doc, '项目涉及的气象数据和算法成果属于内部资料，项目组成员需签署保密协议。代码仓库设置访问权限控制，敏感数据加密存储和传输。', indent=True)

    # ========== 第三部分：项目整体测试方案 ==========
    new_doc.add_heading('项目整体测试方案', level=1)

    new_doc.add_heading('概述', level=2)
    add_para(new_doc, '本测试方案覆盖对流云识别与外推算系统的各个层面，包括数据预处理测试、模型功能测试、性能测试、集成测试和验收测试。测试目标是验证系统各项功能和技术指标满足设计要求。', indent=True)

    new_doc.add_heading('测试资源和环境', level=2)
    new_doc.add_heading('硬件配置', level=3)
    add_para(new_doc, '训练环境：NVIDIA A100/V100 GPU（不少于4卡），CPU 32核以上，内存256GB以上，存储空间10TB以上。推理环境：NVIDIA T4/A10 GPU，CPU 8核以上，内存64GB以上。', indent=True)

    new_doc.add_heading('软件配置', level=3)
    add_para(new_doc, '操作系统：Ubuntu 20.04/Windows Server 2019；深度学习框架：PyTorch 2.0+；CUDA 11.8+；Python 3.9+；依赖库：timm、numpy、h5py、xarray、scikit-learn等。', indent=True)

    new_doc.add_heading('需求分析阶段', level=2)
    add_para(new_doc, '在需求分析阶段，测试团队参与需求评审，确认需求的可测试性，识别测试重点和难点，制定初步测试计划。', indent=True)

    new_doc.add_heading('设计阶段', level=2)
    add_para(new_doc, '在设计阶段，测试团队参与设计评审，了解系统架构和接口设计，编写测试用例，准备测试数据。', indent=True)

    new_doc.add_heading('开发阶段', level=2)
    add_para(new_doc, '开发阶段执行单元测试，开发人员对自己编写的代码进行单元测试，确保各模块功能正确。测试团队同步开发自动化测试脚本。', indent=True)

    new_doc.add_heading('集成测试阶段', level=2)
    add_para(new_doc, '将DIF-UNet识别模块和SatExNet外推模块进行集成测试，验证模块间接口正确性、数据流转正确性和端到端功能完整性。', indent=True)

    new_doc.add_heading('系统测试阶段', level=2)
    add_para(new_doc, '在完整系统环境下进行全面测试，包括功能测试、性能测试、稳定性测试、兼容性测试等，验证系统满足所有需求规格。', indent=True)

    new_doc.add_heading('验收测试', level=2)
    new_doc.add_heading('Alpha测试', level=3)
    add_para(new_doc, '由项目内部测试团队执行，覆盖所有功能点和性能指标，形成Alpha测试报告。', indent=True)

    new_doc.add_heading('Bate测试', level=3)
    add_para(new_doc, '由用户方（气象业务部门）在实际业务环境中执行，使用真实业务数据验证系统的实用性和可靠性。', indent=True)

    new_doc.add_heading('软件测试类型', level=2)
    new_doc.add_heading('白盒测试', level=3)
    add_para(new_doc, '对核心算法模块进行白盒测试，验证代码逻辑的正确性，包括损失函数计算、特征融合操作、自回归推理等关键路径。', indent=True)

    new_doc.add_heading('静态白盒测试', level=3)
    add_para(new_doc, '使用静态代码分析工具（pylint、mypy、bandit）检查代码质量、类型安全和潜在安全漏洞。', indent=True)

    new_doc.add_heading('动态白盒测试', level=3)
    add_para(new_doc, '在运行时进行代码覆盖率和路径覆盖测试，确保关键代码路径均被测试用例覆盖。', indent=True)

    new_doc.add_heading('功能测试', level=2)
    add_para(new_doc, '功能测试覆盖：数据加载与预处理功能、DIF-UNet对流云识别功能、SatExNet外推算功能、结果可视化功能、模型保存与加载功能。每个功能点设计对应的测试用例。', indent=True)

    new_doc.add_heading('UI测试', level=2)
    add_para(new_doc, '如系统提供可视化界面，需测试界面布局合理性、操作便捷性、信息显示准确性。', indent=True)

    new_doc.add_heading('性能测试', level=2)
    new_doc.add_heading('负载测试', level=3)
    add_para(new_doc, '测试系统在满载数据输入下的推理性能，验证单帧推理时间满足业务实时性要求（<30秒/帧）。', indent=True)

    new_doc.add_heading('强度测试', level=3)
    add_para(new_doc, '测试系统在极端输入条件下的稳定性，如输入数据缺失、异常值、超大规模数据等。', indent=True)

    new_doc.add_heading('容量测试', level=3)
    add_para(new_doc, '测试系统处理大规模数据的能力，验证批量处理效率和内存使用情况。', indent=True)

    new_doc.add_heading('安全性和访问控制测试', level=2)
    new_doc.add_heading('应用程序级别的安全性', level=3)
    add_para(new_doc, '测试模型文件和数据的访问权限控制，防止未授权访问。', indent=True)

    new_doc.add_heading('系统级别的安全性', level=3)
    add_para(new_doc, '测试服务器安全配置、网络安全防护、数据加密传输等。', indent=True)

    new_doc.add_heading('故障转移和恢复测试', level=2)
    add_para(new_doc, '测试系统在GPU故障、网络中断等异常情况下的故障恢复能力，验证模型检查点机制确保训练不丢失。', indent=True)

    new_doc.add_heading('兼容性测试', level=2)
    new_doc.add_heading('浏览器兼容性', level=3)
    add_para(new_doc, '如系统提供Web界面，测试在Chrome、Firefox、Edge等主流浏览器上的兼容性。', indent=True)

    new_doc.add_heading('操作系统兼容性', level=3)
    add_para(new_doc, '验证系统在Linux（Ubuntu 20.04）和Windows Server环境下的兼容性。', indent=True)

    new_doc.add_heading('硬件兼容性', level=3)
    add_para(new_doc, '验证系统在不同GPU型号（A100、V100、T4）上的兼容性。', indent=True)

    new_doc.add_heading('安装测试', level=2)
    add_para(new_doc, '测试系统安装部署流程的便捷性和正确性，验证安装文档的完整性。', indent=True)

    new_doc.add_heading('分辨率测试', level=2)
    new_doc.add_heading('说明书测试', level=3)
    add_para(new_doc, '验证用户手册、操作说明等文档的准确性和完整性。', indent=True)

    new_doc.add_heading('宣传材料测试', level=3)
    add_para(new_doc, '验证项目成果展示材料的准确性。', indent=True)

    new_doc.add_heading('帮助文件测试', level=3)
    add_para(new_doc, '验证系统帮助文档的可用性。', indent=True)

    new_doc.add_heading('系统用语', level=3)
    add_para(new_doc, '验证系统输出信息、日志信息、错误提示的准确性和规范性。', indent=True)

    new_doc.add_heading('文档审核测试', level=3)
    add_para(new_doc, '对所有交付文档进行质量审核，确保格式规范、内容准确、版本正确。', indent=True)

    new_doc.add_heading('缺陷管理', level=2)
    new_doc.add_heading('错误跟踪管理系统', level=3)
    add_para(new_doc, '使用GitLab Issue或Jira进行缺陷跟踪管理，每个缺陷有唯一编号、优先级、责任人、状态。', indent=True)

    new_doc.add_heading('软件错误的状态', level=3)
    add_para(new_doc, '缺陷状态包括：新建（New）、已确认（Confirmed）、修复中（In Progress）、已修复（Fixed）、已验证（Verified）、已关闭（Closed）。', indent=True)

    new_doc.add_heading('Bug管理的一般流程', level=3)
    add_para(new_doc, '发现缺陷→提交缺陷报告→项目经理分配→开发人员修复→测试人员验证→关闭缺陷。', indent=True)

    new_doc.add_heading('软件错误流程管理要点', level=3)
    add_para(new_doc, '严重缺陷必须当天修复；一般缺陷3个工作日内修复；缺陷修复后必须经过回归测试验证。', indent=True)

    new_doc.add_heading('环境', level=2)
    add_para(new_doc, '测试环境应与生产环境保持一致，包括操作系统、GPU驱动、CUDA版本、Python版本和依赖库版本。建立独立的测试数据集，不与训练数据重叠。', indent=True)

    # ========== 第四部分：项目实施工作方式 ==========
    new_doc.add_heading('项目实施工作方式', level=1)

    new_doc.add_heading('沟通方式', level=2)
    add_para(new_doc, '项目组内部采用每日站会（15分钟）、每周例会（1小时）制度；与用户方采用双周汇报、里程碑评审会议制度。日常沟通使用即时通讯工具，技术讨论使用邮件或文档协作平台。', indent=True)

    new_doc.add_heading('工作流程', level=2)
    add_para(new_doc, '采用敏捷开发与传统瀑布模型相结合的混合开发模式。算法研发采用迭代式开发，每两周一个迭代周期；系统集成为瀑布模式，按阶段推进。', indent=True)

    # ========== 第五部分：项目培训 ==========
    new_doc.add_heading('项目培训', level=1)

    new_doc.add_heading('培训总体目标', level=2)
    add_para(new_doc, '通过系统培训，使用户方人员掌握对流云识别与外推算系统的使用方法，了解深度学习模型的基本原理，能够独立进行系统操作、结果解读和日常维护。', indent=True)

    new_doc.add_heading('用户培训的定义', level=2)
    new_doc.add_heading('项目管理人员培训', level=3)
    add_para(new_doc, '培训内容包括：系统管理、用户权限配置、数据管理、系统监控等。', indent=True)

    new_doc.add_heading('关键用户培训', level=3)
    add_para(new_doc, '培训内容包括：系统操作使用、结果产品解读、基本故障排查、模型更新流程等。', indent=True)

    new_doc.add_heading('培训策略', level=3)
    add_para(new_doc, '采用理论讲解与实操演练相结合的培训方式，提供详细的操作手册和视频教程。', indent=True)

    new_doc.add_heading('培训方式', level=3)
    add_para(new_doc, '现场培训与远程培训相结合。现场培训不少于2次，每次不少于2天；远程培训提供在线答疑支持。', indent=True)

    new_doc.add_heading('培训设施', level=3)
    add_para(new_doc, '培训场地由用户方提供，配备投影仪、计算机等必要设施。培训环境需安装系统运行所需的软件和硬件。', indent=True)

    new_doc.add_heading('培训计划', level=2)
    add_para(new_doc, '培训安排在项目交付前1个月开始，分为基础培训（系统操作）和进阶培训（模型原理与维护）两个阶段。', indent=True)

    # ========== 第六部分：项目的交付与验收 ==========
    new_doc.add_heading('项目的交付与验收', level=1)

    new_doc.add_heading('集成阶段', level=2)
    add_para(new_doc, '完成DIF-UNet识别模块和SatExNet外推模块的集成，形成完整的对流云识别与外推算系统。集成内容包括：数据输入接口、模型推理引擎、结果输出接口、可视化展示模块。', indent=True)

    new_doc.add_heading('上线阶段', level=2)
    add_para(new_doc, '系统部署至用户方指定的服务器环境，进行上线前最终验证。上线检查清单包括：环境配置确认、数据连接验证、模型加载验证、推理功能验证、性能指标确认。', indent=True)

    new_doc.add_heading('上线后支持阶段', level=2)
    add_para(new_doc, '系统上线后提供3个月的运行保障期，期间提供技术支持、故障排除和必要的参数调优服务。', indent=True)

    # ========== 第七部分：维护及技术服务 ==========
    new_doc.add_heading('维护及技术服务', level=1)

    new_doc.add_heading('服务与支持', level=2)
    add_para(new_doc, '提供7×24小时技术支持热线，工作日8小时内响应，非工作日24小时内响应。', indent=True)

    new_doc.add_heading('服务理念', level=3)
    add_para(new_doc, '以用户需求为导向，以系统稳定运行为目标，提供及时、专业、高效的技术服务。', indent=True)

    new_doc.add_heading('服务组织', level=3)
    add_para(new_doc, '设立专门的技术支持团队，包括项目经理、算法工程师和运维工程师，确保各类技术问题得到及时有效的解决。', indent=True)

    new_doc.add_heading('服务管理', level=3)
    new_doc.add_heading('维护流程', level=4)
    add_para(new_doc, '用户提交服务请求→服务台登记→工程师接单→问题诊断→解决方案→实施修复→用户确认→服务关闭。', indent=True)

    new_doc.add_heading('工作汇报制度', level=4)
    add_para(new_doc, '每月提交系统运行报告，包括系统运行状态、故障统计、性能指标、优化建议等内容。', indent=True)

    new_doc.add_heading('持续改进提升服务质量', level=4)
    add_para(new_doc, '定期收集用户反馈，分析系统运行数据，持续优化模型性能和系统稳定性。每季度进行一次系统评估和优化。', indent=True)

    new_doc.add_heading('维护及技术服务承诺', level=2)
    add_para(new_doc, '项目验收后提供1年免费维护服务，维护期内提供系统升级、Bug修复、性能优化等服务。维护期满后，可根据用户需求签订续保协议。', indent=True)

    # 保存文档
    new_doc.save(OUTPUT)
    print(f"实施方案文档已生成：{OUTPUT}")


if __name__ == '__main__':
    main()
