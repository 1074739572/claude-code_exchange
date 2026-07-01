from docx import Document
import os

templates = [
    "files/样例/5 CWRF区域气候模式重点区域和流域尺度实施方案.docx",
    "files/样例/12  CWRF区域气候模式重点区域和流域尺度技术报告.docx",
    "files/样例/13 CWRF区域气候模式重点区域和流域尺度总结报告.docx"
]

output_file = "template_structures.txt"
with open(output_file, 'w', encoding='utf-8') as f:
    for template_path in templates:
        f.write(f"\n{'='*80}\n")
        f.write(f"Template: {os.path.basename(template_path)}\n")
        f.write('='*80 + '\n')
        
        doc = Document(template_path)
        
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level = int(level)
                except:
                    level = 0
                headings.append((level, para.text))
        
        if headings:
            for level, text in headings:
                indent = "  " * (level - 1) if level > 0 else ""
                f.write(f"{indent}[H{level}] {text}\n")
        else:
            f.write("No headings found. Showing first 20 paragraphs:\n")
            for i, para in enumerate(doc.paragraphs[:20]):
                if para.text.strip():
                    f.write(f"[{i}] {para.text[:100]}\n")

print(f"Template structures extracted to: {output_file}")
