import sys
import io

# Force UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document('files/基于深度学习的对流云识别与外推算法研究_终稿_v1.docx')

with open('_docx_dump.txt', 'w', encoding='utf-8') as f:
    f.write(f'Total paragraphs: {len(doc.paragraphs)}\n')
    count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            count += 1
            if count <= 150:
                f.write(f'[{i}] {p.style.name}: {text[:300]}\n')
    f.write(f'Total non-empty: {count}\n')
    f.write(f'Total tables: {len(doc.tables)}\n')
    
    # Also dump table contents
    for ti, table in enumerate(doc.tables):
        f.write(f'\n--- TABLE {ti} ---\n')
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip()[:50] for c in row.cells]
            f.write(f'  Row {ri}: {" | ".join(cells)}\n')

print('DONE')
