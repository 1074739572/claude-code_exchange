"""Document chunk representation and splitting helpers."""

from __future__ import annotations

import hashlib
import re
from dataclasses import asdict, dataclass
from pathlib import Path

from harness.rag.config import (
    MAX_CHUNK_CHARS,
    MAX_PARENT_CHARS,
    MIN_CHUNK_CHARS,
    OVERLAP_CHARS,
    TARGET_CHUNK_CHARS,
)

HEADING_STYLE_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
}

MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")

LEVEL_PARENT = "parent"
LEVEL_CHILD = "child"


@dataclass
class Chunk:
    id: str
    text: str
    source: str
    heading_path: str
    chapter: str
    section: str
    style: str
    char_count: int
    is_caption: bool = False
    chunk_index: int = 0
    level: str = LEVEL_CHILD
    parent_id: str | None = None
    child_index: int = 0
    modality: str = "text"
    asset_uri: str = ""
    page: int = 0
    table_markdown: str = ""
    image_caption: str = ""
    derived_text: str = ""

    def to_dict(self) -> dict:
        return asdict(self)

    def to_metadata(self) -> dict:
        return {
            "source": self.source,
            "heading_path": self.heading_path,
            "chapter": self.chapter,
            "section": self.section,
            "style": self.style,
            "char_count": self.char_count,
            "is_caption": self.is_caption,
            "chunk_index": self.chunk_index,
            "level": self.level,
            "parent_id": self.parent_id or "",
            "child_index": self.child_index,
            "modality": self.modality,
            "asset_uri": self.asset_uri,
            "page": self.page,
        }


def _make_id(source: str, heading_path: str, chunk_index: int, text: str) -> str:
    digest = hashlib.sha1(
        f"{source}|{heading_path}|{chunk_index}|{text[:64]}".encode()
    ).hexdigest()
    return digest[:16]


def _make_parent_id(source: str, heading_path: str) -> str:
    digest = hashlib.sha1(f"parent|{source}|{heading_path}".encode()).hexdigest()
    return f"p_{digest[:14]}"


def _heading_path(stack: list[tuple[int, str]]) -> str:
    return " > ".join(title for _, title in stack) if stack else "(root)"


def _chapter_section(stack: list[tuple[int, str]]) -> tuple[str, str]:
    if not stack:
        return "", ""
    chapter = stack[0][1]
    section = stack[1][1] if len(stack) > 1 else ""
    return chapter, section


def _format_display(heading: str, body: str) -> str:
    if heading and heading != "(root)":
        return f"{heading}\n\n{body}"
    return body


def _media_chunk(
    *,
    source: str,
    heading: str,
    page: int,
    chunk_index: int,
    modality: str,
    body: str,
    asset_uri: str = "",
    table_markdown: str = "",
    image_caption: str = "",
    derived_text: str = "",
) -> Chunk:
    """Create an independently searchable image or table chunk."""
    chapter = f"第 {page} 页" if page else ""
    return Chunk(
        id=_make_id(source, heading, chunk_index, body),
        text=body,
        source=source,
        heading_path=heading,
        chapter=chapter,
        section="",
        style=f"pdf-{modality}",
        char_count=len(body),
        chunk_index=chunk_index,
        level=LEVEL_CHILD,
        modality=modality,
        asset_uri=asset_uri,
        page=page,
        table_markdown=table_markdown,
        image_caption=image_caption,
        derived_text=derived_text,
    )


def _split_long_text(text: str) -> list[str]:
    if len(text) <= MAX_CHUNK_CHARS:
        return [text] if text.strip() else []

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text]

    pieces: list[str] = []
    current = ""
    for para in paragraphs:
        if len(para) > MAX_CHUNK_CHARS:
            if current:
                pieces.append(current)
                current = ""
            start = 0
            while start < len(para):
                end = min(start + TARGET_CHUNK_CHARS, len(para))
                pieces.append(para[start:end])
                if end >= len(para):
                    break
                start = max(end - OVERLAP_CHARS, start + 1)
            continue
        candidate = f"{current}\n\n{para}".strip() if current else para
        if len(candidate) <= TARGET_CHUNK_CHARS:
            current = candidate
        else:
            if current:
                pieces.append(current)
            current = para
    if current:
        pieces.append(current)
    return pieces


def _truncate_parent_body(body: str) -> str:
    if len(body) <= MAX_PARENT_CHARS:
        return body
    return body[: MAX_PARENT_CHARS - 20] + "\n…[section truncated]"


def _emit_chunks(
    *,
    source: str,
    stack: list[tuple[int, str]],
    body: str,
    style: str,
    is_caption: bool,
    start_index: int,
) -> tuple[list[Chunk], int]:
    heading = _heading_path(stack)
    chapter, section = _chapter_section(stack)
    chunks: list[Chunk] = []
    index = start_index

    merged = body.strip()
    if not merged:
        return chunks, index
    if len(merged) < MIN_CHUNK_CHARS and not is_caption:
        return chunks, index

    if is_caption:
        display = _format_display(heading, merged)
        chunks.append(
            Chunk(
                id=_make_id(source, heading, index, merged),
                text=display,
                source=source,
                heading_path=heading,
                chapter=chapter,
                section=section,
                style=style,
                char_count=len(merged),
                is_caption=True,
                chunk_index=index,
                level=LEVEL_CHILD,
                parent_id=None,
                child_index=0,
            )
        )
        return chunks, index + 1

    parent_id = _make_parent_id(source, heading)
    parent_body = _truncate_parent_body(merged)
    parent_display = _format_display(heading, parent_body)
    chunks.append(
        Chunk(
            id=parent_id,
            text=parent_display,
            source=source,
            heading_path=heading,
            chapter=chapter,
            section=section,
            style=style,
            char_count=len(parent_body),
            is_caption=False,
            chunk_index=index,
            level=LEVEL_PARENT,
            parent_id=None,
            child_index=-1,
        )
    )

    child_pieces = _split_long_text(merged)
    child_index = 0
    for piece in child_pieces:
        if len(piece) < MIN_CHUNK_CHARS:
            continue
        display = _format_display(heading, piece)
        chunks.append(
            Chunk(
                id=_make_id(source, heading, index + child_index, piece),
                text=display,
                source=source,
                heading_path=heading,
                chapter=chapter,
                section=section,
                style=style,
                char_count=len(piece),
                is_caption=False,
                chunk_index=index + child_index,
                level=LEVEL_CHILD,
                parent_id=parent_id,
                child_index=child_index,
            )
        )
        child_index += 1

    if child_index == 0:
        display = _format_display(heading, merged)
        chunks.append(
            Chunk(
                id=_make_id(source, heading, index, merged),
                text=display,
                source=source,
                heading_path=heading,
                chapter=chapter,
                section=section,
                style=style,
                char_count=len(merged),
                is_caption=False,
                chunk_index=index,
                level=LEVEL_CHILD,
                parent_id=parent_id,
                child_index=0,
            )
        )
        child_index = 1

    return chunks, index + child_index


def chunk_markdown(source: str, text: str) -> list[Chunk]:
    stack: list[tuple[int, str]] = []
    body_lines: list[str] = []
    chunks: list[Chunk] = []
    chunk_index = 0

    def flush_body() -> None:
        nonlocal chunk_index, body_lines
        if not body_lines:
            return
        new_chunks, chunk_index = _emit_chunks(
            source=source,
            stack=stack,
            body="\n".join(body_lines),
            style="markdown",
            is_caption=False,
            start_index=chunk_index,
        )
        chunks.extend(new_chunks)
        body_lines = []

    for line in text.splitlines():
        match = MD_HEADING_RE.match(line.strip())
        if match:
            flush_body()
            level = len(match.group(1))
            title = match.group(2).strip()
            stack = [(lvl, name) for lvl, name in stack if lvl < level]
            stack.append((level, title))
            continue
        if line.strip():
            body_lines.append(line.rstrip())
    flush_body()
    return merge_short_chunks(chunks)


def chunk_docx(source: str, path) -> list[Chunk]:
    from docx import Document

    doc = Document(path)
    stack: list[tuple[int, str]] = []
    body_lines: list[str] = []
    chunks: list[Chunk] = []
    chunk_index = 0

    def flush_body(style: str = "Body Text") -> None:
        nonlocal chunk_index, body_lines
        if not body_lines:
            return
        new_chunks, chunk_index = _emit_chunks(
            source=source,
            stack=stack,
            body="\n".join(body_lines),
            style=style,
            is_caption=False,
            start_index=chunk_index,
        )
        chunks.extend(new_chunks)
        body_lines = []

    skip_prefixes = ("编 号", "版 本", "文档控制")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if any(text.startswith(prefix) for prefix in skip_prefixes):
            continue

        style = para.style.name if para.style else "Normal"
        if style in HEADING_STYLE_LEVELS:
            flush_body()
            level = HEADING_STYLE_LEVELS[style]
            stack = [(lvl, name) for lvl, name in stack if lvl < level]
            stack.append((level, text))
            continue

        if style == "Caption":
            flush_body()
            caption_chunks, chunk_index = _emit_chunks(
                source=source,
                stack=stack,
                body=text,
                style="Caption",
                is_caption=True,
                start_index=chunk_index,
            )
            chunks.extend(caption_chunks)
            continue

        if style in ("Body Text", "Normal", "List Paragraph"):
            body_lines.append(text)
            continue

        if style in ("Header", "Footer"):
            continue

        body_lines.append(text)

    flush_body()
    return merge_short_chunks(chunks)


def _markdown_table_parts(rows: list[list[object]]) -> list[str]:
    """Serialize extracted PDF cells into bounded, retrieval-friendly Markdown."""
    cleaned = [
        [re.sub(r"\s+", " ", str(cell or "")).replace("|", "\\|").strip() for cell in row]
        for row in rows
        if any(str(cell or "").strip() for cell in row)
    ]
    if not cleaned:
        return []

    width = max(len(row) for row in cleaned)
    normalized = [row + [""] * (width - len(row)) for row in cleaned]
    header = normalized[0]
    prefix = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    parts: list[str] = []
    current = list(prefix)
    for row in normalized[1:] or [[]]:
        line = "| " + " | ".join(row) + " |"
        if len("\n".join(current + [line])) > MAX_CHUNK_CHARS and len(current) > len(prefix):
            parts.append("\n".join(current))
            current = list(prefix)
        current.append(line)
    if len(current) > len(prefix):
        parts.append("\n".join(current))
    return parts or ["\n".join(prefix)]


def _page_caption(page_text: str) -> str:
    lines = [line.strip() for line in page_text.splitlines() if line.strip()]
    candidates = [
        line for line in lines if re.match(r"^(图|figure|fig\.?)\s*\d", line, re.IGNORECASE)
    ]
    return candidates[0] if candidates else ""


def _vector_regions(page, *, max_regions: int = 4) -> list:
    """Merge substantial vector drawing bounds into VLM-renderable regions."""
    page_area = max(float(page.rect.width * page.rect.height), 1.0)
    regions = []
    for drawing in page.get_drawings():
        rect = drawing.get("rect")
        if not rect or rect.is_empty or rect.width * rect.height < page_area * 0.002:
            continue
        for index, region in enumerate(regions):
            if region.intersects(rect):
                regions[index] = region | rect
                break
        else:
            regions.append(rect)
    regions.sort(key=lambda rect: rect.width * rect.height, reverse=True)
    return regions[:max_regions]


def _is_full_page_image(page, xref: int) -> bool:
    """Detect scanner-created raster images that cover nearly the whole page."""
    try:
        rects = page.get_image_rects(xref)
    except (RuntimeError, ValueError):
        return False
    page_area = max(float(page.rect.width * page.rect.height), 1.0)
    return any(float(rect.width * rect.height) / page_area >= 0.8 for rect in rects)


def chunk_pdf(source: str, path, *, asset_dir) -> list[Chunk]:
    """Extract PDF text, detected tables, and embedded images into RAG chunks."""
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError(
            "PDF ingestion requires PyMuPDF. Install dependencies with: pip install pymupdf"
        ) from exc

    from harness.rag.assets import save_asset
    from harness.rag.enrichment import (
        configured_vlm_model,
        describe_image,
        ocr_dpi,
        ocr_language,
        pdf_ocr_mode,
        should_render_pdf_pages,
        transcribe_scanned_page,
    )

    document = fitz.open(path)
    chunks: list[Chunk] = []
    chunk_index = 0
    seen_images: set[tuple[int, int]] = set()
    try:
        for page_number, page in enumerate(document, start=1):
            page_text = page.get_text("text").strip()
            text_style = "pdf-text"
            mode = pdf_ocr_mode()
            if not page_text and mode in ("vlm", "hybrid"):
                rendered_uri = save_asset(
                    source,
                    name=f"page-{page_number:04d}-ocr-render.png",
                    payload=page.get_pixmap(
                        dpi=ocr_dpi(),
                        alpha=False,
                    ).tobytes("png"),
                    asset_dir=asset_dir,
                )
                try:
                    page_text = transcribe_scanned_page(
                        asset_dir / Path(rendered_uri).name
                    )
                    if page_text:
                        text_style = "pdf-vlm-ocr"
                except Exception:
                    page_text = ""
            if not page_text and mode in ("tesseract", "hybrid"):
                try:
                    text_page = page.get_textpage_ocr(
                        full=True,
                        language=ocr_language(),
                        dpi=ocr_dpi(),
                    )
                    page_text = page.get_text("text", textpage=text_page).strip()
                    text_style = "pdf-tesseract-ocr"
                except Exception:
                    # OCR dependencies are optional; a bad setup must not block
                    # image/table indexing of the rest of the document.
                    page_text = ""
            stack = [(1, f"第 {page_number} 页")]
            if page_text:
                text_chunks, chunk_index = _emit_chunks(
                    source=source,
                    stack=stack,
                    body=page_text,
                    style=text_style,
                    is_caption=False,
                    start_index=chunk_index,
                )
                for chunk in text_chunks:
                    chunk.page = page_number
                chunks.extend(text_chunks)

            try:
                tables = page.find_tables()
                extracted_tables = list(tables.tables)
            except (AttributeError, RuntimeError, ValueError):
                extracted_tables = []
            for table_number, table in enumerate(extracted_tables, start=1):
                for part_number, markdown in enumerate(_markdown_table_parts(table.extract()), start=1):
                    title = f"[表格] PDF 第 {page_number} 页，表 {table_number}"
                    if part_number > 1:
                        title += f"，分段 {part_number}"
                    body = f"{title}\n\n{markdown}"
                    chunks.append(
                        _media_chunk(
                            source=source,
                            heading=f"第 {page_number} 页 > 表 {table_number}",
                            page=page_number,
                            chunk_index=chunk_index,
                            modality="table",
                            body=body,
                            table_markdown=markdown,
                            derived_text=body,
                        )
                    )
                    chunk_index += 1

            caption = _page_caption(page_text)
            page_images = page.get_images(full=True)
            # Charts in generated PDFs are commonly vector drawings. Detect and
            # render their regions even on mixed pages that also contain logos
            # or photos, so those raster assets do not suppress chart indexing.
            for region_number, region in enumerate(
                _vector_regions(page)
                if configured_vlm_model() and should_render_pdf_pages()
                else [],
                start=1,
            ):
                rendered_uri = save_asset(
                    source,
                    name=f"page-{page_number:04d}-vector-region-{region_number:02d}.png",
                    payload=page.get_pixmap(
                        matrix=fitz.Matrix(2, 2), clip=region, alpha=False
                    ).tobytes("png"),
                    asset_dir=asset_dir,
                )
                rendered_path = asset_dir / Path(rendered_uri).name
                try:
                    rendered_description = describe_image(rendered_path, surrounding_text=page_text)
                except Exception:
                    rendered_description = None
                if rendered_description:
                    body = (
                        f"[图片] PDF 第 {page_number} 页，矢量图形区域 {region_number}\n"
                        f"页面文本提示：{page_text[:1200] or '无可提取文本'}\n"
                        f"视觉描述：{rendered_description}"
                    )
                    chunks.append(
                        _media_chunk(
                            source=source,
                            heading=f"第 {page_number} 页 > 矢量图形 {region_number}",
                            page=page_number,
                            chunk_index=chunk_index,
                            modality="image",
                            body=body,
                            asset_uri=rendered_uri,
                            derived_text=rendered_description,
                        )
                    )
                    chunk_index += 1

            for image_number, image in enumerate(page_images, start=1):
                xref = int(image[0])
                if text_style == "pdf-vlm-ocr" and _is_full_page_image(page, xref):
                    # The rendered page was already sent to the VLM for faithful
                    # transcription; avoid a second paid description call.
                    continue
                identity = (page_number, xref)
                if identity in seen_images:
                    continue
                seen_images.add(identity)
                try:
                    extracted = document.extract_image(xref)
                except RuntimeError:
                    continue
                payload = extracted.get("image")
                if not payload:
                    continue
                extension = extracted.get("ext", "png")
                asset_uri = save_asset(
                    source,
                    name=f"page-{page_number:04d}-image-{image_number:02d}.{extension}",
                    payload=payload,
                    asset_dir=asset_dir,
                )
                asset_path = asset_dir / Path(asset_uri).name
                try:
                    vision_description = describe_image(asset_path, surrounding_text=page_text)
                except Exception:
                    vision_description = None
                fallback = (
                    "未生成视觉描述：请设置 HARNESS_RAG_VLM_MODEL 和 "
                    "HARNESS_RAG_VLM_API_KEY（或 OPENAI_API_KEY）后重新索引。"
                )
                description = vision_description or fallback
                body = (
                    f"[图片] PDF 第 {page_number} 页，图 {image_number}\n"
                    f"图注：{caption or '未检测到'}\n"
                    f"页面文本提示：{page_text[:1200] or '无可提取文本'}\n"
                    f"视觉描述：{description}"
                )
                chunks.append(
                    _media_chunk(
                        source=source,
                        heading=f"第 {page_number} 页 > 图 {image_number}",
                        page=page_number,
                        chunk_index=chunk_index,
                        modality="image",
                        body=body,
                        asset_uri=asset_uri,
                        image_caption=caption,
                        derived_text=description,
                    )
                )
                chunk_index += 1
    finally:
        document.close()
    return merge_short_chunks(chunks)


def merge_short_chunks(chunks: list[Chunk]) -> list[Chunk]:
    """Merge adjacent short child chunks under the same parent."""
    if not chunks:
        return chunks

    parents = [chunk for chunk in chunks if chunk.level == LEVEL_PARENT]
    children = [chunk for chunk in chunks if chunk.level != LEVEL_PARENT]
    merged_children = _merge_short_children(children)
    return parents + merged_children


def _merge_short_children(children: list[Chunk]) -> list[Chunk]:
    if not children:
        return children

    merged: list[Chunk] = []
    buffer: Chunk | None = None

    def flush_buffer() -> None:
        nonlocal buffer
        if buffer is not None:
            merged.append(buffer)
            buffer = None

    for chunk in children:
        if chunk.is_caption or chunk.modality != "text" or not chunk.parent_id:
            flush_buffer()
            merged.append(chunk)
            continue

        if buffer is None:
            if len(chunk.text) < MIN_CHUNK_CHARS:
                buffer = chunk
            else:
                merged.append(chunk)
            continue

        same_parent = (
            buffer.parent_id == chunk.parent_id
            and buffer.heading_path == chunk.heading_path
            and buffer.source == chunk.source
        )
        combined_len = len(buffer.text) + len(chunk.text)
        if same_parent and combined_len <= MAX_CHUNK_CHARS:
            buffer = _join_chunks(buffer, chunk)
            if len(buffer.text) >= MIN_CHUNK_CHARS:
                flush_buffer()
            continue

        flush_buffer()
        if len(chunk.text) < MIN_CHUNK_CHARS:
            buffer = chunk
        else:
            merged.append(chunk)

    flush_buffer()
    return merged


def merge_short_chunk_dicts(chunks: list[dict]) -> list[dict]:
    if not chunks:
        return chunks
    objects = [
        Chunk(
            id=item["id"],
            text=item["text"],
            source=item["source"],
            heading_path=item["heading_path"],
            chapter=item["chapter"],
            section=item["section"],
            style=item["style"],
            char_count=item["char_count"],
            is_caption=bool(item.get("is_caption")),
            chunk_index=int(item.get("chunk_index", 0)),
            level=item.get("level", LEVEL_CHILD),
            parent_id=item.get("parent_id"),
            child_index=int(item.get("child_index", 0)),
            modality=item.get("modality", "text"),
            asset_uri=item.get("asset_uri", ""),
            page=int(item.get("page", 0)),
            table_markdown=item.get("table_markdown", ""),
            image_caption=item.get("image_caption", ""),
            derived_text=item.get("derived_text", ""),
        )
        for item in chunks
    ]
    return [chunk.to_dict() for chunk in merge_short_chunks(objects)]


def _join_chunks(left: Chunk, right: Chunk) -> Chunk:
    body_left = left.text.split("\n\n", 1)[-1] if "\n\n" in left.text else left.text
    body_right = right.text.split("\n\n", 1)[-1] if "\n\n" in right.text else right.text
    combined_body = f"{body_left}\n\n{body_right}".strip()
    heading = left.heading_path
    display = _format_display(heading, combined_body)
    return Chunk(
        id=left.id,
        text=display,
        source=left.source,
        heading_path=left.heading_path,
        chapter=left.chapter,
        section=left.section,
        style=left.style,
        char_count=len(combined_body),
        is_caption=False,
        chunk_index=left.chunk_index,
        level=LEVEL_CHILD,
        parent_id=left.parent_id,
        child_index=left.child_index,
        modality=left.modality,
        asset_uri=left.asset_uri,
        page=left.page,
        table_markdown=left.table_markdown,
        image_caption=left.image_caption,
        derived_text=left.derived_text,
    )
